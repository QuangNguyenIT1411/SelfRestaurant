from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path("/mnt/c/Users/Quang/Downloads/SINH_VIEN/SINH_VIEN/SelfRestaurant-main")
SRC = ROOT / "BaiTapSo3_ArchitectureViews"


def set_default_font(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        style = "Heading 1"
    elif level == 2:
        style = "Heading 2"
    else:
        style = "Heading 3"
    p.style = style
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def add_diagram_placeholder(doc: Document, title: str, source_txt: str) -> None:
    add_paragraph(doc, f"Biểu đồ: {title}", bold=True)
    add_paragraph(doc, f"Nguồn mã Mermaid: {source_txt}", italic=True)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = True
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        paragraph.add_run("\n")
    run = paragraph.add_run(f"[ CHÈN HÌNH {title.upper()} TẠI ĐÂY ]")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    for _ in range(6):
        paragraph.add_run("\n")
    doc.add_paragraph()


def parse_markdown(md_path: Path):
    items = []
    for raw in md_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            items.append(("blank", ""))
            continue
        if line.startswith("### "):
            items.append(("h3", line[4:]))
        elif line.startswith("## "):
            items.append(("h2", line[3:]))
        elif line.startswith("# "):
            items.append(("h1", line[2:]))
        elif line.startswith("- "):
            items.append(("bullet", line[2:]))
        else:
            items.append(("p", line))
    return items


def fill_from_markdown(doc: Document, md_path: Path, placeholder_rules: dict[str, tuple[str, str]]) -> None:
    items = parse_markdown(md_path)
    for kind, text in items:
        if kind == "h1":
            add_title(doc, text)
            doc.add_paragraph()
        elif kind == "h2":
            add_heading(doc, text, 1)
            rule = placeholder_rules.get(text)
            if rule:
                add_diagram_placeholder(doc, rule[0], rule[1])
        elif kind == "h3":
            add_heading(doc, text, 2)
        elif kind == "bullet":
            add_bullet(doc, text)
        elif kind == "p":
            add_paragraph(doc, text)


def build_bai1() -> Document:
    doc = Document()
    set_default_font(doc)
    fill_from_markdown(
        doc,
        SRC / "Bai1_Logical_View.md",
        {
            "2. Giới thiệu Use Case tổng thể": ("Biểu đồ Use Case tổng thể", "UseCase_Diagram.txt"),
            "3. Giới thiệu các lớp trong biểu đồ lớp": ("Biểu đồ lớp tổng thể", "Class_Diagram.txt"),
        },
    )
    return doc


def build_bai2() -> Document:
    doc = Document()
    set_default_font(doc)
    fill_from_markdown(
        doc,
        SRC / "Bai2_Implementation_View.md",
        {
            "2. Giới thiệu các package": ("Biểu đồ gói tổng thể", "Package_Diagram.txt"),
            "3. Giới thiệu các component": ("Biểu đồ thành phần tổng thể", "Component_Diagram.txt"),
        },
    )
    return doc


def build_bai3() -> Document:
    doc = Document()
    set_default_font(doc)
    fill_from_markdown(
        doc,
        SRC / "Bai3_Process_View.md",
        {
            "2. Giải thích biểu đồ hoạt động 1: Khách hàng chọn bàn và đặt món": (
                "Biểu đồ hoạt động 1: Khách hàng chọn bàn và đặt món",
                "Activity_Customer_Order.txt",
            ),
            "3. Giải thích biểu đồ hoạt động 2: Bếp xử lý món và khách xác nhận nhận món": (
                "Biểu đồ hoạt động 2: Bếp xử lý món và khách xác nhận nhận món",
                "Activity_Chef_Fulfillment.txt",
            ),
            "4. Giải thích biểu đồ hoạt động 3: Thu ngân thanh toán QR và cộng điểm": (
                "Biểu đồ hoạt động 3: Thu ngân thanh toán QR và cộng điểm",
                "Activity_Cashier_QR_Payment.txt",
            ),
        },
    )
    return doc


def build_bai4() -> Document:
    doc = Document()
    set_default_font(doc)
    fill_from_markdown(
        doc,
        SRC / "Bai4_Deployment_View.md",
        {
            "2. Giới thiệu các thành phần triển khai": ("Biểu đồ triển khai tổng thể", "Deployment_Diagram.txt"),
        },
    )
    return doc


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def finalize(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    add_page_number(section)


def save_doc(name: str, builder) -> Path:
    doc = builder()
    finalize(doc)
    out = SRC / name
    doc.save(out)
    return out


def main() -> None:
    outputs = [
        save_doc("Bai1_Logical_View.docx", build_bai1),
        save_doc("Bai2_Implementation_View.docx", build_bai2),
        save_doc("Bai3_Process_View.docx", build_bai3),
        save_doc("Bai4_Deployment_View.docx", build_bai4),
    ]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
